import tkinter as tk
from tkinter import filedialog as fd
from tkinter import messagebox as mb
from docx import Document  # pip install python-docx
from docx.enum.section import WD_ORIENT
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Pt, Mm
import os.path
import math
from make_df import *


class Application(tk.Tk):
    def __init__(self):
        super().__init__()
        self.title("Подсчет диагностики")
        self.resizable(0, 0)
        self.main_frame = tk.Frame(self)
        self.type_frame = tk.Frame(self.main_frame)
        self.type_lbl = tk.Label(self.main_frame, text="тип диагностики")
        self.diag_type = tk.StringVar()
        self.diag_type.set("первичная")
        self.diag_type_radio_0 = tk.Radiobutton(
            text="первичная",
            variable=self.diag_type,
            value="первичной",
            master=self.type_frame,
        )
        self.diag_type_radio_1 = tk.Radiobutton(
            text="промежуточная",
            variable=self.diag_type,
            value="промежуточной",
            master=self.type_frame
        )
        self.diag_type_radio_2 = tk.Radiobutton(
            text="итоговая",
            variable=self.diag_type,
            value="итоговой",
            master=self.type_frame
        )

        self.date_lbl = tk.Label(self.main_frame, text="учебный год, например 2021-2022")
        self.date_ent = tk.Entry(self.main_frame, width=60)

        self.teacher_lbl = tk.Label(self.main_frame, text="ФИО педагога")
        self.teacher_ent = tk.Entry(self.main_frame)

        self.discipline_lbl = tk.Label(
            self.main_frame,
            text="название программы (без кавычек)"
        )
        self.discipline_ent = tk.Entry(self.main_frame)

        self.level_lbl = tk.Label(self.main_frame, text="уровень программы")
        self.level_ent = tk.Entry(self.main_frame)

        self.year_lbl = tk.Label(self.main_frame, text="год обучения (прописью)")
        self.year_ent = tk.Entry(self.main_frame)

        self.group_lbl = tk.Label(self.main_frame, text="номер или название группы")
        self.group_ent = tk.Entry(self.main_frame)

        self.source_file_path_lbl = tk.Label(
            self.main_frame,
            text="файл с оценками не выбран",
            fg="#ff0000"
        )
        self.choose_source_file_path_btn = tk.Button(
            self.main_frame, text="выбрать файл с оценками"
        )
        self.choose_source_file_path_btn["command"] = self.get_file_path

        self.save_btn = tk.Button(
            self.main_frame, text="сохранить результат диагностики"
        )

        # размещаем виджеты по сетке
        self.main_frame.grid(padx="20", pady="20")
        self.type_lbl.grid(row=0, column=0, sticky='e')
        self.type_frame.grid(row=0, column=1)
        self.diag_type_radio_0.grid(row=0, column=0)
        self.diag_type_radio_1.grid(row=0, column=1)
        self.diag_type_radio_2.grid(row=0, column=2)

        self.date_lbl.grid(row=4, column=0, sticky='e', pady="10")
        self.date_ent.grid(row=4, column=1, columnspan=2, sticky='nesw', padx="5", pady="10")

        self.teacher_lbl.grid(row=5, column=0, sticky='e', pady="10")
        self.teacher_ent.grid(row=5, column=1, columnspan=2, sticky='nesw', padx="5", pady="10")

        self.discipline_lbl.grid(row=6, column=0, sticky='e', pady="10")
        self.discipline_ent.grid(row=6, column=1, columnspan=2, sticky='nesw', padx="5", pady="10")

        self.level_lbl.grid(row=7, column=0, sticky='e', pady="10")
        self.level_ent.grid(row=7, column=1, columnspan=2, sticky='nesw', padx="5", pady="10")

        self.year_lbl.grid(row=8, column=0, sticky='e', pady="10")
        self.year_ent.grid(row=8, column=1, columnspan=2, sticky='nesw', padx="5", pady="10")

        self.group_lbl.grid(row=9, column=0, sticky='e', pady="10")
        self.group_ent.grid(row=9, column=1, columnspan=2, sticky='nesw', padx="5", pady="10")

        self.source_file_path_lbl.grid(row=10, column=0, columnspan=3, sticky='nesw', pady="20")
        self.choose_source_file_path_btn.grid(row=11, column=0, columnspan=3)
        self.save_btn.grid(row=12, column=0, columnspan=3, pady="10")
        self.save_btn["command"] = self.save_result_to_file

        self.save_btn["state"] = "disabled"

    def get_file_path(self):
        self.source_file_path = fd.askopenfilename(
            filetypes=[("excel file", ".xlsx")]
        )
        self.source_file_path_lbl.config(text=self.source_file_path)

        if os.path.isfile(self.source_file_path):
            self.save_btn["state"] = "normal"
            self.source_file_path_lbl["fg"] = "#196f3d"
        else:
            self.save_btn["state"] = "disabled"

    def save_result_to_file(self):
        try:
            self.students_df = make_students_df(self.source_file_path)
        except Exception:
            print(self.students_df)
            self.save_btn["state"] = "disabled"
            self.source_file_path_lbl["fg"] = "#ff0000"
            mb.showwarning("Ошибка!", f"Невозможно обработать оценки, выбирите другой файл!")
        else:
            # TODO: разделить на создание таблицы и сохранение файла
            path = fd.asksaveasfile(
                mode="w",
                defaultextension='.docx',
                filetypes=[("word file", ".docx")]
            )

            document = Document()

            styles = document.styles

            styles.add_style('my_diagnostics', WD_STYLE_TYPE.PARAGRAPH)
            style = document.styles['my_diagnostics']
            font = style.font
            font.name = 'Times New Roman'
            font.size = Pt(12)

            styles.add_style('my_table', WD_STYLE_TYPE.TABLE)
            style = document.styles['my_table']
            font = style.font
            font.name = 'Times New Roman'
            font.size = Pt(12)
            font.italic = True

            styles['Table Grid'].font.name = font.name = 'Times New Roman'
            styles['Table Grid'].font.size = Pt(12)

            # делаем заголовок
            p = document.add_paragraph(style="my_diagnostics")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run("ПРОТОКОЛ")
            p.add_run(
                f"\nподведения итогов {self.diag_type.get()} "
                f"диагностики дополнительной общеразвивающей программы\n"
                f"{self.date_ent.get()} уч. год"
            )

            p = document.add_paragraph(style="my_diagnostics")
            p.add_run(
                f"Педагог ДО {self.teacher_ent.get()}\n"
                f"Общеразвивающая  программа «{self.discipline_ent.get()}»\n"
                f"Уровень программы, год обучения, номер (название) группы: "
                f"{self.level_ent.get()}, {self.year_ent.get()}, "
                f"{self.group_ent.get()}"
            )

            # делаем шапку
            table = document.add_table(rows=1, cols=15)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.style = "Table Grid"
            table.allow_autofit = True
            table.autofit = True
            heading_cells = table.rows[0].cells

            # соединяем ячейки
            table.cell(0, 2).merge(table.cell(0, 4))
            table.cell(0, 5).merge(table.cell(0, 7))
            table.cell(0, 8).merge(table.cell(0, 10))
            table.cell(0, 12).merge(table.cell(0, 14))

            # заполняем шапку
            heading_cells[0].text = "№ пп"
            heading_cells[1].text = "Ф.И. обучающегося"
            heading_cells[2].text = "Предметные результаты (П)"
            heading_cells[5].text = "Метапредметные результаты (М)"
            heading_cells[8].text = "Личностные результаты (Л)"
            heading_cells[11].text = "В баллах (П+М+Л):3"
            heading_cells[12].text = "Уровень начальной подготовки"
            "по дополнительной общеразвивающей программе"

            for num, student in self.students_df.iterrows():
                cells = table.add_row().cells
                table.cell(num, 2).merge(table.cell(num, 4))
                table.cell(num, 5).merge(table.cell(num, 7))
                table.cell(num, 8).merge(table.cell(num, 10))
                table.cell(num, 12).merge(table.cell(num, 14))
                cells[0].text = str(num)
                cells[1].text = student['ФИО']
                cells[2].text = student["Предметный уровень"]
                cells[5].text = student["Метапредметный уровень"]
                cells[8].text = student["Личностный уровень"]
                cells[11].text = str(round(student["В баллах (П + М + Л) / 3"], 2))
                cells[12].text = student["Уровень начальной подготовки"]

            cells = table.add_row().cells
            table.cell(-1, 0).merge(table.cell(-1, 1))
            cells[0].text = "уровень освоения"
            cells[1].text = ""
            cells[2].text = "Н"
            cells[3].text = "С"
            cells[4].text = "В"
            cells[5].text = "Н"
            cells[6].text = "С"
            cells[7].text = "В"
            cells[8].text = "Н"
            cells[9].text = "С"
            cells[10].text = "В"
            cells[11].text = ""
            cells[12].text = "Н"
            cells[13].text = "С"
            cells[14].text = "В"

            cells = table.add_row().cells
            table.cell(-1, 0).merge(table.cell(-1, 1))
            cells[0].text = "количество учащихся"
            cells[1].text = ""

            # счет предметных
            filtr = self.students_df["Предметный уровень"] == "Н"
            cells[2].text = str(filtr.sum())
            filtr = self.students_df["Предметный уровень"] == "С"
            cells[3].text = str(filtr.sum())
            filtr = self.students_df["Предметный уровень"] == "В"
            cells[4].text = str(filtr.sum())

            # счет метапредметных
            filtr = self.students_df["Метапредметный уровень"] == "Н"
            cells[5].text = str(filtr.sum())
            filtr = self.students_df["Метапредметный уровень"] == "С"
            cells[6].text = str(filtr.sum())
            filtr = self.students_df["Метапредметный уровень"] == "В"
            cells[7].text = str(filtr.sum())

            # счет личностных
            filtr = self.students_df["Личностный уровень"] == "Н"
            cells[8].text = str(filtr.sum())
            filtr = self.students_df["Личностный уровень"] == "С"
            cells[9].text = str(filtr.sum())
            filtr = self.students_df["Личностный уровень"] == "В"
            cells[10].text = str(filtr.sum())

            # счет начальной подготовки
            cells[11].text = ""
            filtr = self.students_df["Уровень начальной подготовки"] == "Н"
            cells[12].text = str(filtr.sum())
            filtr = self.students_df["Уровень начальной подготовки"] == "С"
            cells[13].text = str(filtr.sum())
            filtr = self.students_df["Уровень начальной подготовки"] == "В"
            cells[14].text = str(filtr.sum())

            cells = table.add_row().cells
            table.cell(-1, 0).merge(table.cell(-1, 1))
            cells[0].text = "% соотношение"

            # % предметных
            filtr = self.students_df["Предметный уровень"] == "Н"
            cells[2].text = str(round(filtr.sum() * 100 / len(self.students_df.index), 2))
            filtr = self.students_df["Предметный уровень"] == "С"
            cells[3].text = str(round(filtr.sum() * 100 / len(self.students_df.index), 2))
            filtr = self.students_df["Предметный уровень"] == "В"
            cells[4].text = str(round(filtr.sum() * 100 / len(self.students_df.index), 2))

            # % мета
            filtr = self.students_df["Метапредметный уровень"] == "Н"
            cells[5].text = str(round(filtr.sum() * 100 / len(self.students_df.index), 2))
            filtr = self.students_df["Метапредметный уровень"] == "С"
            cells[6].text = str(round(filtr.sum() * 100 / len(self.students_df.index), 2))
            filtr = self.students_df["Метапредметный уровень"] == "В"
            cells[7].text = str(round(filtr.sum() * 100 / len(self.students_df.index), 2))

            # % личностных
            filtr = self.students_df["Личностный уровень"] == "Н"
            cells[8].text = str(round(filtr.sum() * 100 / len(self.students_df.index), 2))
            filtr = self.students_df["Личностный уровень"] == "С"
            cells[9].text = str(round(filtr.sum() * 100 / len(self.students_df.index), 2))
            filtr = self.students_df["Личностный уровень"] == "В"
            cells[10].text = str(round(filtr.sum() * 100 / len(self.students_df.index), 2))

            # % начальной подготовки
            filtr = self.students_df["Уровень начальной подготовки"] == "Н"
            cells[12].text = str(round(filtr.sum() * 100 / len(self.students_df.index), 2))
            filtr = self.students_df["Уровень начальной подготовки"] == "С"
            cells[13].text = str(round(filtr.sum() * 100 / len(self.students_df.index), 2))
            filtr = self.students_df["Уровень начальной подготовки"] == "В"
            cells[14].text = str(round(filtr.sum() * 100 / len(self.students_df.index), 2))

            # альбомная ориентация
            section = document.sections[-1]
            section.orientation = WD_ORIENT.LANDSCAPE
            new_width, new_height = section.page_height, section.page_width
            section.orientation = WD_ORIENT.LANDSCAPE
            section.page_width = new_width
            section.page_height = new_height
            # поля
            section.top_margin = Mm(13)
            section.bottom_margin = Mm(13)
            section.left_margin = Mm(30)
            section.right_margin = Mm(13)

            for row in table.rows:
                row.cells[0].width = Mm(10)
                row.cells[1].width = Mm(60)
                row.cells[11].width = Mm(20)

            document.save(path.name)
            mb.showinfo("Сохранено!", f"Диагностика успешно сохранена в файл {path.name}")

if __name__ == "__main__":
    window = Application()
    window.mainloop()
